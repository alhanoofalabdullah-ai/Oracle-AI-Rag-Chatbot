from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
}


@dataclass
class LoadedDocument:
    """
    Normalized document representation.
    """

    text: str
    file_name: str
    metadata: dict


def _load_pdf(
    data: bytes,
) -> tuple[str, dict]:

    reader = PdfReader(
        BytesIO(data)
    )

    pages: list[str] = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):

        text = (
            page.extract_text()
            or ""
        )

        if text.strip():

            pages.append(
                f"[Page {page_number}]\n"
                f"{text.strip()}"
            )

    return (
        "\n\n".join(pages),
        {
            "page_count": len(
                reader.pages
            )
        },
    )


def _load_docx(
    data: bytes,
) -> tuple[str, dict]:

    document = DocxDocument(
        BytesIO(data)
    )

    paragraphs: list[str] = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:

            paragraphs.append(text)

    return (
        "\n\n".join(paragraphs),
        {
            "paragraph_count": len(
                paragraphs
            )
        },
    )


def _load_text(
    data: bytes,
) -> tuple[str, dict]:

    text = data.decode(
        "utf-8",
        errors="replace",
    )

    return (
        text,
        {},
    )


def load_document(
    file_name: str,
    data: bytes,
) -> LoadedDocument:

    extension = (
        Path(file_name)
        .suffix
        .lower()
    )

    if extension not in SUPPORTED_EXTENSIONS:

        raise ValueError(
            f"Unsupported file type: "
            f"{extension}. "
            f"Supported types: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if extension == ".pdf":

        text, metadata = _load_pdf(data)

    elif extension == ".docx":

        text, metadata = _load_docx(data)

    else:

        text, metadata = _load_text(data)

    if not text.strip():

        raise ValueError(
            f"No readable text found "
            f"in '{file_name}'."
        )

    metadata.update(
        {
            "file_name": file_name,
            "extension": extension,
        }
    )

    return LoadedDocument(
        text=text.strip(),
        file_name=file_name,
        metadata=metadata,
    )
